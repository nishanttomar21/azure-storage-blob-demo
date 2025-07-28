import os
from azure.identity import DefaultAzureCredential
from azure.storage.blob import BlobServiceClient


# =============================================================================
# STEP 1: CREATE SAMPLE FILES FOR TESTING
# =============================================================================

def create_sample_files():
    """Create sample files if they don't exist in the current directory"""
    # Create sample1.txt if it doesn't exist
    if not os.path.exists('sample1.txt'):
        with open('sample1.txt', 'w') as f:
            f.write("This is a sample text file for Azure Blob Storage testing.")
        print("Created sample1.txt")

    # Create sample2.docx if it doesn't exist (requires python-docx library)
    if not os.path.exists('sample2.docx'):
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Azure Blob Demo', 0)
            doc.add_paragraph('This is a sample Word document for blob upload testing.')
            doc.save('sample2.docx')
            print("Created sample2.docx\n")
        except ImportError:
            print("Warning: python-docx not installed. Creating sample2.txt instead.")
            with open('sample2.txt', 'w') as f:
                f.write("Sample document content (as .txt since docx library not available)")


# =============================================================================
# STEP 2: CONFIGURE AZURE STORAGE CONNECTION
# =============================================================================

def configure_azure_connection(account_url):
    """Configure and return Azure Blob Service Client"""
    print("Authenticating with Azure...")

    # Check if placeholder URL is being used
    if '<your_storage_account_name>' in account_url:
        raise ValueError("Please replace <your_storage_account_name> with your actual account name.")

    credential = DefaultAzureCredential()
    blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)

    return blob_service_client


# =============================================================================
# STEP 3: CREATE BLOB CONTAINER
# =============================================================================

def create_blob_container(blob_service_client, container_name):
    """Create a blob container if it doesn't exist"""
    container_client = blob_service_client.get_container_client(container_name)

    try:
        container_client.create_container()
        print(f"Container '{container_name}' created successfully!")
    except Exception as e:
        print(f"Container may already exist: {e}")

    return container_client


# =============================================================================
# STEP 4: UPLOAD FILES TO BLOB STORAGE
# =============================================================================

def upload_files_to_blob(container_client, upload_files):
    """Upload multiple files to blob storage"""
    print(f"\nStarting upload of {len(upload_files)} files...")

    successful_uploads = []

    for local_file in upload_files:
        try:
            # Check if the file exists locally before attempting upload
            if not os.path.exists(local_file):
                print(f"WARNING: File '{local_file}' not found. Skipping.")
                continue

            # Open file in binary mode and upload to blob storage
            with open(local_file, 'rb') as data:
                blob_client = container_client.get_blob_client(local_file)
                blob_client.upload_blob(data, overwrite=True)
                print(f"Uploaded: {local_file}")
                successful_uploads.append(local_file)

        except Exception as e:
            print(f"Error uploading {local_file}: {e}")

    return successful_uploads


# =============================================================================
# STEP 5: LIST ALL BLOBS IN THE CONTAINER
# =============================================================================

def list_blobs_in_container(container_client, container_name):
    """List all blobs in the container and return their names"""
    print(f"\nListing all blobs in container '{container_name}':")

    blob_names = []

    try:
        blobs = container_client.list_blobs()
        blob_count = 0

        for blob in blobs:
            print(f"   - {blob.name}")
            blob_names.append(blob.name)
            blob_count += 1

        if blob_count == 0:
            print("   (No blobs found in container)")
        else:
            print(f"   Total blobs: {blob_count}")

    except Exception as e:
        print(f"Error listing blobs: {e}")

    return blob_names


# =============================================================================
# STEP 6: DOWNLOAD BLOBS FROM STORAGE
# =============================================================================

def download_blobs_from_storage(container_client, upload_files):
    """Download blobs from storage to local files with 'downloaded_' prefix"""
    print(f"\nStarting download process...")

    downloaded_files = []

    for src_blob_name in upload_files:
        try:
            # Define output filename with "downloaded_" prefix
            downloaded_filename = "downloaded_" + src_blob_name

            # Get blob client for the specific blob
            blob_client = container_client.get_blob_client(src_blob_name)

            # Download the blob to local file
            with open(downloaded_filename, "wb") as download_file:
                download_stream = blob_client.download_blob()
                download_file.write(download_stream.readall())
                print(f"Downloaded '{src_blob_name}' as '{downloaded_filename}'")
                downloaded_files.append(downloaded_filename)

        except Exception as e:
            print(f"Error processing {src_blob_name}: {e}")

    return downloaded_files


# =============================================================================
# STEP 7: DELETE THE CONTAINER
# =============================================================================

def delete_container(container_client, container_name):
    """Delete the blob container and all its contents"""
    print(f"\nDeleting container '{container_name}'...")

    try:
        container_client.delete_container()
        print(f"Container '{container_name}' deleted successfully!")
        return True
    except Exception as e:
        print(f"Error deleting container: {e}")
        return False


# =============================================================================
# STEP 8: FINAL VERIFICATION
# =============================================================================

def verify_local_files(expected_files):
    """Verify the existence of expected local files"""
    print(f"\nLocal files in directory:")

    file_status = {}

    for file in expected_files:
        if os.path.exists(file):
            print(f"   [FOUND] {file}")
            file_status[file] = True
        else:
            print(f"   [MISSING] {file}")
            file_status[file] = False

    return file_status


# =============================================================================
# MAIN FUNCTION TO ORCHESTRATE ALL STEPS
# =============================================================================

def main():
    """Main function that orchestrates all Azure Blob Storage operations"""
    # Configuration
    account_url = "https://<your_storage_account_name>.blob.core.windows.net/"
    container_name = 'mytestcontainer'
    upload_files = ['sample1.txt', 'sample2.docx']
    expected_files = ['sample1.txt', 'sample2.docx', 'downloaded_sample1.txt', 'downloaded_sample2.docx']

    try:
        # Step 1: Create sample files
        create_sample_files()

        # Step 2: Configure Azure connection
        blob_service_client = configure_azure_connection(account_url)

        # Step 3: Create blob container
        container_client = create_blob_container(blob_service_client, container_name)

        # Step 4: Upload files to blob storage
        successful_uploads = upload_files_to_blob(container_client, upload_files)

        # Step 5: List all blobs in container
        blob_names = list_blobs_in_container(container_client, container_name)

        # Step 6: Download blobs from storage
        downloaded_files = download_blobs_from_storage(container_client, successful_uploads)

        # Step 7: Delete the container
        deletion_success = delete_container(container_client, container_name)

        # Step 8: Final verification
        file_status = verify_local_files(expected_files)

        # Return summary information
        return {
            'successful_uploads': successful_uploads,
            'blob_names': blob_names,
            'downloaded_files': downloaded_files,
            'deletion_success': deletion_success,
            'file_status': file_status
        }

    except Exception as e:
        print(f"Script failed with error: {e}")
        return None


# =============================================================================
# SCRIPT ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    result = main()

    if result:
        print(f"\nSUMMARY:")
        print(f"- Files uploaded: {len(result['successful_uploads'])}")
        print(f"- Files downloaded: {len(result['downloaded_files'])}")
        print(f"- Container deleted: {result['deletion_success']}")
        print(f"\nScript execution completed!")